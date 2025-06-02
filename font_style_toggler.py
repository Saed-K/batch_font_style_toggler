import sys
import random
from pathlib import Path

import spacy
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QDialog, QFormLayout, QDialogButtonBox,
    QSpinBox, QComboBox, QMessageBox, QFileDialog, QListWidget, QListWidgetItem,
    QProgressBar, QTextBrowser, QLineEdit
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from docx import Document
from docx.shared import RGBColor

# --- NLP setup ---
nlp = spacy.load("en_core_web_sm")

# --- Rule definition ---
class StyleRule:
    def __init__(self, target, action, percent, extra=None):
        self.target = target
        self.action = action
        self.percent = percent
        self.extra = extra

    def description(self):
        desc = f"{self.action} {self.percent}% of {self.target}"
        if self.extra:
            desc += f" ({self.extra})"
        return desc

# --- Background worker for batch processing ---
class BatchWorker(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal()

    def __init__(self, files, rules, output_dir):
        super().__init__()
        self.files = files
        self.rules = rules
        self.output_dir = output_dir

    def run(self):
        total = len(self.files)
        for i, filepath in enumerate(self.files, start=1):
            try:
                processor = DocumentProcessor(Path(filepath), self.rules, self.output_dir)
                processor.apply()
            except Exception:
                pass
            self.progress.emit(int(i / total * 100))
        self.finished.emit()

# --- Processor for .docx and .md ---
class DocumentProcessor:
    def __init__(self, filepath: Path, rules: list[StyleRule], output_dir: Path):
        self.filepath = filepath
        self.rules = rules
        self.output_dir = output_dir

    def apply(self):
        if self.filepath.suffix.lower() == '.docx':
            self._apply_docx()
        else:
            self._apply_md()

    def _apply_docx(self):
        doc = Document(self.filepath)
        # Apply heading rules
        for rule in [r for r in self.rules if r.target == 'heading']:
            self._apply_heading_rule(doc, rule)
        # Apply other rules
        text_rules = [r for r in self.rules if r.target != 'heading']
        if text_rules:
            for para in doc.paragraphs:
                self._apply_text_rules(para, text_rules)
        # Save
        self.output_dir.mkdir(parents=True, exist_ok=True)
        out_path = self.output_dir / f"{self.filepath.stem}_styled.docx"
        doc.save(out_path)

    def _apply_heading_rule(self, doc, rule):
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        count = max(1, int(len(headings) * rule.percent / 100))
        for p in random.sample(headings, count):
            run = p.runs[0] if p.runs else p.add_run(p.text)
            self._style_run(run, rule)

    def _apply_text_rules(self, para, rules):
        tokens = nlp(para.text)
        style_map = {i: [] for i in range(len(tokens))}
        for rule in rules:
            idxs = [i for i, tok in enumerate(tokens) if tok.pos_.lower() == rule.target]
            chosen = random.sample(idxs, max(1, int(len(idxs) * rule.percent / 100))) if idxs else []
            for i in chosen:
                style_map[i].append(rule)
        # Rebuild paragraph
        para._element.clear()
        for i, tok in enumerate(tokens):
            run = para.add_run(tok.text_with_ws)
            for rule in style_map[i]:
                self._style_run(run, rule)

    def _style_run(self, run, rule: StyleRule):
        if rule.action == 'bold':
            run.bold = True
        elif rule.action == 'italic':
            run.italic = True
        elif rule.action == 'underline':
            run.underline = True
        elif rule.action == 'strikethrough':
            run.font.strike = True
        elif rule.action == 'uppercase':
            run.text = run.text.upper()
        elif rule.action == 'color' and rule.extra:
            r, g, b = rule.extra
            run.font.color.rgb = RGBColor(r, g, b)

    def _apply_md(self):
        # Placeholder: similar logic for Markdown using **, *, etc.
        text = self.filepath.read_text(encoding='utf-8')
        # ... implement Markdown rules ...
        self.output_dir.mkdir(parents=True, exist_ok=True)
        out_path = self.output_dir / f"{self.filepath.stem}_styled.md"
        out_path.write_text(text, encoding='utf-8')

# --- Custom ListWidget supporting Delete key ---
class DeletableListWidget(QListWidget):
    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Delete:
            for item in self.selectedItems():
                self.takeItem(self.row(item))
        else:
            super().keyPressEvent(e)

# --- Main GUI ---
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Batch Font-Style Toggler")
        self.resize(1000, 600)

        container = QWidget()
        self.setCentralWidget(container)
        main_layout = QHBoxLayout(container)

        # File panel
        self.fileList = DeletableListWidget()
        add_files_btn = QPushButton("Add Files...")
        add_files_btn.clicked.connect(self.add_files)
        remove_files_btn = QPushButton("Remove Selected Files")
        remove_files_btn.clicked.connect(self.remove_files)
        file_layout = QVBoxLayout()
        file_layout.addWidget(QLabel("Documents:"))
        file_layout.addWidget(self.fileList)
        file_layout.addWidget(add_files_btn)
        file_layout.addWidget(remove_files_btn)

        # Output directory
        self.outDir = QLineEdit(str(Path.cwd()))
        browse_out_btn = QPushButton("Browse Out Dir")
        browse_out_btn.clicked.connect(self.browse_output)
        file_layout.addWidget(QLabel("Output Folder:"))
        file_layout.addWidget(self.outDir)
        file_layout.addWidget(browse_out_btn)

        # Rules panel
        self.rulesList = DeletableListWidget()
        add_rule_btn = QPushButton("Add Rule")
        add_rule_btn.clicked.connect(self.new_rule)
        remove_rule_btn = QPushButton("Remove Selected Rules")
        remove_rule_btn.clicked.connect(self.remove_rules)
        rule_layout = QVBoxLayout()
        rule_layout.addWidget(QLabel("Style Rules:"))
        rule_layout.addWidget(self.rulesList)
        rule_layout.addWidget(add_rule_btn)
        rule_layout.addWidget(remove_rule_btn)

        # Preview
        self.preview = QTextBrowser()
        self.fileList.currentItemChanged.connect(self.update_preview)

        # Footer
        self.progress = QProgressBar()
        run_btn = QPushButton("Run Batch")
        run_btn.clicked.connect(self.start_batch)
        footer_layout = QHBoxLayout()
        footer_layout.addWidget(self.progress)
        footer_layout.addWidget(run_btn)

        # Assemble
        main_layout.addLayout(file_layout, 2)
        main_layout.addLayout(rule_layout, 2)
        main_layout.addWidget(self.preview, 3)
        main_layout.addLayout(footer_layout, 1)

    def add_files(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "Select Documents", "", "Docs (*.docx *.md)"
        )
        for p in paths:
            self.fileList.addItem(p)

    def remove_files(self):
        for item in self.fileList.selectedItems():
            self.fileList.takeItem(self.fileList.row(item))

    def browse_output(self):
        d = QFileDialog.getExistingDirectory(self, "Select Output Directory")
        if d:
            self.outDir.setText(d)

    def new_rule(self):
        dlg = RuleDialog(self)
        if dlg.exec_():
            rule = dlg.get_rule()
            item = QListWidgetItem(rule.description())
            item.setData(Qt.UserRole, rule)
            self.rulesList.addItem(item)

    def remove_rules(self):
        for item in self.rulesList.selectedItems():
            self.rulesList.takeItem(self.rulesList.row(item))

    def update_preview(self, current, previous=None):
        if not current:
            self.preview.clear()
            return
        path = Path(current.text())
        if path.suffix.lower() == '.md':
            text = path.read_text(encoding='utf-8')
            self.preview.setPlainText(text)
        else:
            try:
                doc = Document(path)
                text = '\n'.join(p.text for p in doc.paragraphs)
                self.preview.setPlainText(text)
            except Exception:
                self.preview.setPlainText('Cannot preview this file.')

    def start_batch(self):
        files = [self.fileList.item(i).text() for i in range(self.fileList.count())]
        rules = [self.rulesList.item(i).data(Qt.UserRole) for i in range(self.rulesList.count())]
        output_dir = Path(self.outDir.text())
        self.worker = BatchWorker(files, rules, output_dir)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.finished.connect(lambda: QMessageBox.information(self, "Done", "Batch complete!"))
        self.worker.start()

# --- Rule Dialog ---
class RuleDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Create New Style Rule")
        layout = QFormLayout(self)

        self.targetBox = QComboBox()
        self.targetBox.addItems(['heading', 'verb', 'adjective', 'noun', 'adverb'])
        self.actionBox = QComboBox()
        self.actionBox.addItems([ 'bold', 'italic', 'underline', 'strikethrough', 'uppercase', 'color' ])
        self.percentSpin = QSpinBox()
        self.percentSpin.setRange(0, 100)
        self.percentSpin.setValue(100)
        self.colorPicker = QLineEdit('255,0,0')
        self.colorPicker.setEnabled(False)
        self.actionBox.currentTextChanged.connect(lambda t: self.colorPicker.setEnabled(t == 'color'))

        layout.addRow("Target:", self.targetBox)
        layout.addRow("Action:", self.actionBox)
        layout.addRow("Percentage:", self.percentSpin)
        layout.addRow("Color RGB:", self.colorPicker)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addRow(buttons)

    def get_rule(self) -> StyleRule:
        extra = None
        if self.actionBox.currentText() == 'color':
            rgb = tuple(map(int, self.colorPicker.text().split(',')))
            extra = rgb
        return StyleRule(
            target=self.targetBox.currentText(),
            action=self.actionBox.currentText(),
            percent=self.percentSpin.value(),
            extra=extra
        )

# --- Entry point ---
def main():
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
